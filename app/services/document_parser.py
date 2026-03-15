import logging
from io import BytesIO
from typing import Optional

logger = logging.getLogger(__name__)

class DocumentParser:

    SUPPORTED_TYPES = {
        "application/pdf" : "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain" : "txt",
    }

    def parse(self, file_bytes: bytes, content_type : str, filename: str = "") -> str:
        doc_type = self.SUPPORTED_TYPES.get(content_type)

        if not doc_type and filename:
            if filename.endswith(".pdf"):
                doc_type = "pdf"
            elif filename.endswith(".docx"):
                doc_type = "docx"
            elif filename.endswith(".txt"):
                doc_type = "txt"


        if not doc_type:
            raise ValueError(
                f"unsupported file type: {content_type}. "
                f"supported PDF, DOCX, TXT"
            )

        logger.info(f"Parsing {doc_type.upper()} file: {filename or 'unnames'}")

        if doc_type == "pdf":
            return self._parse_pdf(file_bytes)
        elif doc_type == "docx":
            return self._parse_docx(file_bytes)
        else:
            return self._parse_txt(file_bytes)


    def _parse_pdf(self, file_bytes) -> str:

        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(file_bytes))
            pages = []

            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():

                    pages.append(f"[Page {page_num + 1}]\n{text.strip()}")  

                    if not pages:
                     raise ValueError("PDF contains no extractable text ")

                    full_text = "\n\n".join(pages)
                    logger.info(f"PDF parsed: {len(reader.pages)} pages, {len(full_text)} chars")
                    return full_text

        except ImportError:
            raise ImportError("pypdf not installed. Run: pip install pypdf")
        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            raise

    def _parse_docx(self, file_bytes: bytes) -> str:
        try:
            from docx import Document
            doc = Document(BytesIO(file_bytes))
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)
            if not paragraphs:
                raise ValueError("DOCX contains no extractable text")
            full_text = "\n\n".join(paragraphs)
            logger.info(f"DOCX parsed: {len(paragraphs)} paragraphs, {len(full_text)} chars")
            return full_text
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            raise
    def _parse_txt(self, file_bytes: bytes) -> str:
        """Decode plain text file."""
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                text = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                text = file_bytes.decode("latin-1")
            text = text.strip()
            if not text:
                raise ValueError("TXT file is empty")
            logger.info(f"TXT parsed: {len(text)} chars")
            return text
        except Exception as e:
            logger.error(f"TXT parsing failed: {e}")
            raise
# Singleton — import this everywhere
document_parser = DocumentParser()






